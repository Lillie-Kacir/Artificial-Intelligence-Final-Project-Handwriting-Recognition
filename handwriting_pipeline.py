import argparse
import re
import shutil
from dataclasses import dataclass
from pathlib import Path

import cv2
import numpy as np
import pytesseract
from docx import Document
from spellchecker import SpellChecker

DEFAULT_WINDOWS_TESSERACT_PATHS = (
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
)


@dataclass
class OcrCandidate:
    text: str
    score: float
    average_confidence: float
    strategy: str


def configure_tesseract() -> None:
    if shutil.which("tesseract"):
        return
    for path in DEFAULT_WINDOWS_TESSERACT_PATHS:
        if Path(path).exists():
            pytesseract.pytesseract.tesseract_cmd = path
            return


def deskew(gray_image: np.ndarray) -> np.ndarray:
    _, thresh = cv2.threshold(
        gray_image, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU
    )
    coords = np.column_stack(np.where(thresh > 0))
    if coords.size == 0:
        return gray_image

    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle

    if abs(angle) < 0.25:
        return gray_image

    height, width = gray_image.shape[:2]
    matrix = cv2.getRotationMatrix2D((width // 2, height // 2), angle, 1.0)
    return cv2.warpAffine(
        gray_image,
        matrix,
        (width, height),
        flags=cv2.INTER_CUBIC,
        borderMode=cv2.BORDER_REPLICATE,
    )


def preprocess_image_for_ocr(image_path: Path) -> dict[str, np.ndarray]:
    image = cv2.imread(str(image_path))
    if image is None:
        raise FileNotFoundError(f"Could not read image: {image_path}")

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    upscaled = cv2.resize(gray, None, fx=1.5, fy=1.5, interpolation=cv2.INTER_CUBIC)
    contrast = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8)).apply(upscaled)
    denoised = cv2.GaussianBlur(contrast, (3, 3), 0)
    aligned = deskew(denoised)

    _, otsu = cv2.threshold(aligned, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    adaptive = cv2.adaptiveThreshold(
        aligned,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31,
        11,
    )
    adaptive_inv = cv2.bitwise_not(adaptive)
    return {
        "gray": aligned,
        "otsu": otsu,
        "adaptive": adaptive,
        "adaptive_inv": adaptive_inv,
    }


def score_candidate(text: str, confidences: list[float]) -> float:
    avg_conf = sum(confidences) / len(confidences) if confidences else 0.0
    alpha_chars = sum(ch.isalpha() for ch in text)
    alpha_ratio = alpha_chars / max(1, len(text))
    return avg_conf + (alpha_ratio * 14.0) + min(len(text.split()), 30) * 0.25


def extract_text(preprocessed_variants: dict[str, np.ndarray]) -> tuple[str, OcrCandidate]:
    configure_tesseract()
    candidates: list[OcrCandidate] = []

    for name, image in preprocessed_variants.items():
        for psm in (6, 4, 11):
            config = f"--oem 3 --psm {psm}"
            raw_text = pytesseract.image_to_string(image, config=config).strip()
            if not raw_text:
                continue

            data = pytesseract.image_to_data(
                image,
                config=config,
                output_type=pytesseract.Output.DICT,
            )
            confidences = []
            for conf in data.get("conf", []):
                try:
                    value = float(conf)
                except (TypeError, ValueError):
                    continue
                if value >= 0:
                    confidences.append(value)

            score = score_candidate(raw_text, confidences)
            candidate = OcrCandidate(
                text=raw_text,
                score=score,
                average_confidence=(sum(confidences) / len(confidences)) if confidences else 0.0,
                strategy=f"{name}:psm{psm}",
            )
            candidates.append(candidate)

    if not candidates:
        raise RuntimeError(
            "OCR produced no text. Verify that Tesseract is installed and the image is readable."
        )

    best = max(candidates, key=lambda candidate: candidate.score)
    return best.text, best


def edit_distance(lhs: str, rhs: str) -> int:
    if lhs == rhs:
        return 0
    if not lhs:
        return len(rhs)
    if not rhs:
        return len(lhs)

    previous = list(range(len(rhs) + 1))
    for i, ch_l in enumerate(lhs, start=1):
        current = [i]
        for j, ch_r in enumerate(rhs, start=1):
            cost = 0 if ch_l == ch_r else 1
            current.append(
                min(
                    previous[j] + 1,
                    current[j - 1] + 1,
                    previous[j - 1] + cost,
                )
            )
        previous = current
    return previous[-1]


def normalize_ocr_token(token: str) -> str:
    if any(ch.isalpha() for ch in token) and any(ch.isdigit() for ch in token):
        return token.translate(str.maketrans({"0": "o", "1": "l", "5": "s"}))
    return token


def choose_best_correction(word: str, candidates: set[str], spell: SpellChecker) -> str:
    if not candidates:
        return word

    filtered = list(candidates)
    if len(word) >= 3:
        same_edges = [c for c in filtered if c[0] == word[0] and c[-1] == word[-1]]
        if same_edges:
            filtered = same_edges

    return min(
        filtered,
        key=lambda candidate: (
            edit_distance(word, candidate),
            -spell.word_frequency.dictionary.get(candidate, 0),
            abs(len(candidate) - len(word)),
        ),
    )


def predict_words(raw_text: str) -> str:
    spell = SpellChecker(distance=2)
    token_pattern = re.compile(r"[A-Za-z']+|[^A-Za-z']+")
    tokens = token_pattern.findall(raw_text)

    corrected = []
    for token in tokens:
        if not re.fullmatch(r"[A-Za-z']+", token):
            corrected.append(token)
            continue

        # Preserve very short words and likely acronyms.
        if len(token) <= 2 or token.isupper():
            corrected.append(token)
            continue

        normalized = normalize_ocr_token(token.lower())
        if normalized in spell:
            corrected.append(token)
            continue

        candidate_words = spell.candidates(normalized) or set()
        replacement = choose_best_correction(normalized, candidate_words, spell)

        # Guardrail: avoid aggressive substitutions that change too many characters.
        if edit_distance(normalized, replacement) > max(2, len(normalized) // 2):
            corrected.append(token)
            continue

        if token[0].isupper():
            replacement = replacement.capitalize()
        corrected.append(replacement)

    return "".join(corrected)


def save_outputs(predicted_text: str, output_stem: Path) -> tuple[Path, Path]:
    txt_path = output_stem.with_suffix(".txt")
    docx_path = output_stem.with_suffix(".docx")

    txt_path.parent.mkdir(parents=True, exist_ok=True)
    txt_path.write_text(predicted_text, encoding="utf-8")

    document = Document()
    document.add_heading("Predicted Handwriting Text", level=1)
    document.add_paragraph(predicted_text)
    document.save(str(docx_path))

    return txt_path, docx_path


def run_pipeline(image_path: Path, output_stem: Path) -> tuple[str, str, Path, Path]:
    preprocessed = preprocess_image_for_ocr(image_path)
    raw_text, _ = extract_text(preprocessed)
    predicted_text = predict_words(raw_text)
    txt_path, docx_path = save_outputs(predicted_text, output_stem)
    return raw_text, predicted_text, txt_path, docx_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Read a handwritten image, recognize letters/words, apply "
            "word prediction/correction, and export output documents."
        )
    )
    parser.add_argument(
        "--image",
        required=True,
        help="Path to the input handwriting image file.",
    )
    parser.add_argument(
        "--output-stem",
        default="outputs/predicted_document",
        help="Output file stem (without extension).",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    image_path = Path(args.image)
    output_stem = Path(args.output_stem)

    raw_text, predicted_text, txt_path, docx_path = run_pipeline(image_path, output_stem)

    print("Raw OCR text:")
    print(raw_text.strip())
    print("\nPredicted text:")
    print(predicted_text.strip())
    print(f"\nSaved text output: {txt_path}")
    print(f"Saved document output: {docx_path}")


if __name__ == "__main__":
    main()
